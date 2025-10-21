from docx import Document
from docx.shared import Pt
import json

def generate_full_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)

    # 三、详细设计
    doc.add_heading('三、详细设计', level=1)
    
    # 3.2 数据库设计
    doc.add_heading('3.2 数据库设计', level=2)
    doc.add_heading('3.2.1 问题数据', level=3)
    doc.add_paragraph('''表结构：
- id: 主键
- title: 问题标题
- content: 问题内容
- creator_id: 创建用户
- status: 状态(待解决/已解决)
- created_at: 创建时间''')
    doc.add_heading('3.2.2 用户数据', level=3)
    doc.add_paragraph('''表结构：
- id: 主键
- username: 用户名
- password: 加密密码
- role: 角色(学生/管理员)
- avatar: 头像路径
- created_at: 注册时间''')
    
    # 3.1 界面设计
    doc.add_heading('3.1 界面设计', level=2)
    
    # 3.1.1 管理员界面
    doc.add_heading('3.1.1 管理员界面', level=3)
    doc.add_paragraph('''功能说明：
- 问题管理：审核/删除问题
- 用户管理：封禁/解封用户  
- 数据统计：查看平台使用数据
UI特点：红色警示按钮，数据可视化图表''')

    # 3.1.2 学生界面
    doc.add_heading('3.1.2 学生界面', level=3)
    doc.add_paragraph('''功能说明:
- 问题列表：分类浏览问题
- 提问功能：发布新问题
- 个人中心：查看我的提问/收藏
UI特点：蓝色主色调，简洁卡片布局''')

    # 3.1.3 问题详情界面  
    doc.add_heading('3.1.3 问题详情界面', level=3)
    doc.add_paragraph('''功能说明:
- 问题内容展示
- 回答列表及点赞功能
- 收藏/分享按钮
UI特点：Markdown渲染，交互式评论区''')

    # 3.1.4 其他界面
    doc.add_heading('3.1.4 其他界面', level=3)  
    doc.add_paragraph('''1. 收藏栏：
- 收藏的问题列表
- 分类管理功能

2. 历史记录栏：
- 浏览历史时间线
- 快速跳转功能

3. 个人信息栏（点击头像进入）：
- 个人资料编辑
- 账号安全设置

4. 消息栏：
- 系统通知
- 回答被采纳提醒''')
    
    # 3.3 关键算法
    doc.add_heading('3.3 关键算法', level=2)
    algorithms = [
        '数据库连接',
        '注册登录',
        '问题搜索',
        '问题发布/评论',
        '点赞收藏',
        '问题处理'
    ]
    for i, algo in enumerate(algorithms, 1):
        doc.add_heading(f'3.3.{i} {algo}', level=3)
        doc.add_paragraph('待补充算法流程图和伪代码')

    # 四、测试报告
    doc.add_heading('四、测试报告', level=1)
    tests = [
        '问题发送和处理',
        '点赞收藏评论', 
        '问题搜索',
        '消息功能'
    ]
    for i, test in enumerate(tests, 1):
        doc.add_heading(f'4.{i} {test}测试', level=2)
        doc.add_paragraph('待补充测试用例和结果')

    # 五、使用说明
    doc.add_heading('五、网站使用', level=1)
    doc.add_heading('5.1 环境要求', level=2)
    doc.add_paragraph('Python 3.8+, Node.js 16+, MySQL 5.7+')
    doc.add_heading('5.2 使用说明', level=2)
    doc.add_paragraph('待补充部署步骤和操作指南')

    # 六、项目总结
    doc.add_heading('六、项目总结', level=1)
    doc.add_heading('6.2 项目困难', level=2)
    doc.add_paragraph('待补充技术难点和解决方案')
    doc.add_heading('6.3 开发感悟', level=2)
    doc.add_paragraph('待补充团队协作经验')
    doc.add_heading('6.4 后续安排', level=2)
    doc.add_paragraph('前端改进和微服务嵌入计划')

    import os
    import tempfile
    
    save_paths = [
        '反馈通系统文档.docx',
        os.path.join(os.path.expanduser("~"), 'Desktop', '反馈通系统文档.docx'),
        os.path.join(tempfile.gettempdir(), '反馈通系统文档.docx')
    ]
    
    saved = False
    for path in save_paths:
        try:
            doc.save(path)
            print(f"文档已成功生成到: {path}")
            saved = True
            break
        except Exception as e:
            print(f"尝试保存到 {path} 失败: {str(e)}")
    
    if not saved:
        print("无法保存文档，请检查文件权限或磁盘空间")

if __name__ == '__main__':
    generate_full_doc()